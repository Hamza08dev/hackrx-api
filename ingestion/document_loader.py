#!/usr/bin/env python3
"""
Document Loader Module
- Extracts text from PDF, DOCX, and TXT files
- Handles multiple document formats reliably
- Clean, focused functionality
"""

import os
import logging
from typing import Optional
import pdfplumber
from docx import Document
import PyPDF2

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber (primary) with PyPDF2 fallback."""
    try:
        # Primary method: pdfplumber
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if text.strip():
                logger.info(f"✅ PDF extracted with pdfplumber: {len(text)} chars")
                return text.strip()
    
    except Exception as e:
        logger.warning(f"⚠️ pdfplumber failed: {e}, trying PyPDF2...")
    
    try:
        # Fallback method: PyPDF2
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            logger.info(f"✅ PDF extracted with PyPDF2: {len(text)} chars")
            return text.strip()
    
    except Exception as e:
        logger.error(f"❌ Both PDF extraction methods failed: {e}")
        raise ValueError(f"Failed to extract text from PDF: {e}")

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files."""
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        logger.info(f"✅ DOCX extracted: {len(text)} chars")
        return text.strip()
    
    except Exception as e:
        logger.error(f"❌ DOCX extraction failed: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {e}")

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT files with encoding detection."""
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
                logger.info(f"✅ TXT extracted with {encoding}: {len(text)} chars")
                return text.strip()
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception as e:
            logger.error(f"❌ TXT extraction failed with {encoding}: {e}")
            continue
    
    raise ValueError(f"Failed to extract text from TXT file with any encoding")

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from various file formats.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text as string
        
    Raises:
        ValueError: If file format not supported or extraction fails
    """
    if not os.path.exists(file_path):
        raise ValueError(f"File not found: {file_path}")
    
    # Get file extension
    _, ext = os.path.splitext(file_path.lower())
    
    logger.info(f"📄 Extracting text from {ext} file: {os.path.basename(file_path)}")
    
    try:
        if ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return extract_text_from_docx(file_path)
        elif ext == '.txt':
            return extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    except Exception as e:
        logger.error(f"❌ Text extraction failed for {file_path}: {e}")
        raise

def get_supported_extensions() -> list:
    """Get list of supported file extensions."""
    return ['.pdf', '.docx', '.txt']

def is_supported_file(file_path: str) -> bool:
    """Check if file format is supported."""
    _, ext = os.path.splitext(file_path.lower())
    return ext in get_supported_extensions()